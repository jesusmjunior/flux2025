import streamlit as st
import json
from graphviz import Digraph
from PIL import Image

st.set_page_config(page_title="Fluxograma COGEX", layout="centered")

st.image("cogex.png", width=120)
st.title("COGEX - CORREGEDORIA DO FORO EXTRAJUDICIAL DO ESTADO DO MARANHÃO")
st.markdown("### Sistema de Modelagem de Processos")

st.markdown("---")
arquivo_escolhido = st.selectbox("🔽 Selecione um fluxograma", ["fluxo2.json"])

setores_cogex = [
    "Gabinete dos Juízes Corregedores",
    "Núcleo de Registro Civil",
    "Núcleo de Atos",
    "Diretoria do COGEX",
    "Chefe de Gabinete dos Juízes Corregedores",
    "Assessoria Jurídica",
    "Coordenadoria de Serventias Extrajudiciais",
    "Coordenadoria de Reclamações e Processos Disciplinares",
    "Coordenadoria Administrativa",
    "Coordenadoria de Inspeções",
    "Coordenadoria de Análise de Contas"
]

setor_escolhido = st.selectbox("🏛️ Setor:", setores_cogex)

try:
    with open(arquivo_escolhido, encoding="utf-8") as f:
        dados = json.load(f)
except Exception as e:
    st.error(f"❌ Erro ao carregar ou renderizar o fluxo: {e}")
    st.stop()

st.markdown(f"### 📌 {dados.get('titulo', 'Sem Título')}")
st.markdown(f"**Setor Selecionado:** {setor_escolhido}")

col1, col2 = st.columns([3, 1])

with col1:
    fluxo = Digraph("Fluxograma")
    fluxo.attr(rankdir="TB", size="8,10", nodesep="0.5")

    estilo_map = {
        "inicio": {"shape": "circle", "style": "filled", "fillcolor": "lightgreen"},
        "tarefa": {"shape": "box", "style": "rounded,filled", "fillcolor": "lightblue"},
        "verificacao": {"shape": "box", "style": "rounded,filled", "fillcolor": "khaki"},
        "publicacao": {"shape": "box", "style": "rounded,filled", "fillcolor": "lightpink"},
        "fiscalizacao": {"shape": "box", "style": "rounded,filled", "fillcolor": "lightgrey"},
        "fim": {"shape": "doublecircle", "style": "filled", "fillcolor": "red"}
    }

    for etapa in dados["etapas"]:
        estilo = estilo_map.get(etapa["tipo"], {})
        fluxo.node(etapa["id"], etapa["texto"], **estilo)

    for origem, destino in dados["conexoes"]:
        fluxo.edge(origem, destino)

    st.graphviz_chart(fluxo)

with col2:
    st.subheader("📘 Legenda")
    for tipo, estilo in estilo_map.items():
        simbolo = "⬤" if "circle" in estilo["shape"] else "⬛"
        st.markdown(f"{simbolo} **{tipo.capitalize()}** – cor `{estilo['fillcolor']}`")

    st.markdown("---")
    st.subheader("⚖️ Base Legal")
    st.markdown(dados.get("base_legal", "Não definida"))

# Exportação sem dependência do Graphviz executável
with st.expander("🖨️ Exportar visualização como HTML A4"):
    html_export = f'''
<!DOCTYPE html>
<html lang="pt-br">
<head>
<meta charset="UTF-8">
<title>Exportação Fluxo</title>
<style>
body {{
    font-family: Arial, sans-serif;
    max-width: 800px;
    margin: auto;
    padding: 40px;
}}
.header img {{ height: 80px; }}
.header {{ text-align: center; }}
</style>
</head>
<body>
    <div class="header">
        <h1>CORREGEDORIA DO FORO EXTRAJUDICIAL</h1>
        <h2>{dados.get("titulo", "")}</h2>
        <p><strong>Setor:</strong> {setor_escolhido}</p>
    </div>
    <p><strong>📘 Legenda:</strong><br>
    ⬤ Início – lightgreen<br>
    ⬛ Tarefa – lightblue<br>
    ⬛ Verificação – khaki<br>
    ⬛ Publicação – lightpink<br>
    ⬛ Fiscalização – lightgrey<br>
    ⬤ Fim – red</p>
    <p><strong>⚖️ Base Legal:</strong><br>{dados.get("base_legal", "")}</p>
</body>
</html>
    '''
    st.download_button("📥 Baixar HTML estilo A4", data=html_export, file_name="fluxo_visual.html", mime="text/html")


# === Exportação para DOCX com imagem real ===
from graphviz import Digraph
from docx import Document
from docx.shared import Inches
from io import BytesIO
from PIL import Image

if st.button("📤 Exportar layout para Word (.docx)"):
    fluxo_doc = Digraph("Fluxo")
    fluxo_doc.attr(rankdir="TB", size="8,10", nodesep="0.5")

    for etapa in dados["etapas"]:
        estilo = estilo_map.get(etapa["tipo"], {})
        fluxo_doc.node(etapa["id"], etapa["texto"], **estilo)

    for origem, destino in dados["conexoes"]:
        fluxo_doc.edge(origem, destino)

    try:
        img_bytes = fluxo_doc.pipe(format="png")
        img_io = BytesIO(img_bytes)
        img = Image.open(img_io)

        doc = Document()
        doc.add_heading("CORREGEDORIA DO FORO EXTRAJUDICIAL DO ESTADO DO MARANHÃO", 0)
        doc.add_heading(dados.get("titulo", ""), level=1)
        doc.add_paragraph(f"Setor: {setor_escolhido}")
        img_io.seek(0)
        doc.add_picture(img_io, width=Inches(5.5))

        doc.add_heading("📘 Legenda", level=2)
        doc.add_paragraph("⬤ Início – lightgreen")
        doc.add_paragraph("⬛ Tarefa – lightblue")
        doc.add_paragraph("⬛ Verificação – khaki")
        doc.add_paragraph("⬛ Publicação – lightpink")
        doc.add_paragraph("⬛ Fiscalização – lightgrey")
        doc.add_paragraph("⬤ Fim – red")

        doc.add_heading("⚖️ Base Legal", level=2)
        doc.add_paragraph(dados.get("base_legal", "Não definida"))

        doc.add_paragraph("Sistema de Modelagem de Processos – COGEX/TJMA")

        output_docx = "fluxograma_COGEX.docx"
        doc.save(output_docx)

        with open(output_docx, "rb") as f:
            st.download_button("📥 Baixar Word (.docx)", f, file_name=output_docx, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    except Exception as e:
        st.error(f"Erro na exportação: {e}")
